"""
Build M11408924_林罕宏_問卷調查設計.docx
Research diagram = editable VML shapes, placed AFTER hypothesis table, BEFORE news page
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(__file__)
OUTPUT = os.path.join(BASE_DIR, 'M11408924_林罕宏_問卷調查設計.docx')

doc = Document()

# ---- Page setup ----
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

# ---- Helpers ----
def shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def cell_text(cell, text, bold=False, size=12, color=None):
    cell.text = ''; p = cell.paragraphs[0]
    r = p.add_run(text); r.font.name='Microsoft JhengHei'; r.font.size=Pt(size); r.font.bold=bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    if color: r.font.color.rgb=RGBColor(*color)

def tbl_borders(table):
    table._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'))

def merge(table, row, c1, c2):
    table.cell(row, c1).merge(table.cell(row, c2))

def R(para, text, bold=False, size=12, italic=False, color=None):
    r = para.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.name='Microsoft JhengHei'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    if color: r.font.color.rgb=RGBColor(*color)
    return r


# ======== 表1 ========
p=doc.add_paragraph(); R(p,'表 1　問卷調查設計之程序說明表',bold=True,size=14)

t1=doc.add_table(rows=9,cols=4); t1.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t1)

# Row0
cell_text(t1.cell(0,0),'研究主題:',bold=True,size=11); merge(t1,0,1,3)
cell_text(t1.cell(0,1),'新型態廣告整合 AI Overview 拯救 Google 廣告衝擊研究',size=11); shading(t1.cell(0,1),'FFF2CC')

# Row1
cell_text(t1.cell(1,0),'1 以80字以內，提出<合作請求>文案',bold=True,size=11); merge(t1,1,1,3)
cell_text(t1.cell(1,1),'您好，我是國立臺北科技大學研究生林罕宏。正在進行「AI 摘要對搜尋廣告點擊意願與營收效益之影響」研究，誠摯邀請您撥冗填寫問卷（約 5 分鐘），所有資料僅作學術用途並嚴格保密，感謝您的寶貴協助！（共 78 字）',size=11)
shading(t1.cell(1,1),'FFF2CC')

# Row2
cell_text(t1.cell(2,0),'2 挑選1-2個主要變數',bold=True,size=11); merge(t1,2,1,3)
c=t1.cell(2,1); c.text=''; p=c.paragraphs[0]
R(p,'自變數 IV1：AI Overview Intrusiveness（AI 概覽侵入性）',bold=True,size=11)
p.add_run('\n')
R(p,'自變數 IV2：Innovative Ad Format Integration（創新廣告格式整合）',bold=True,size=11)
p.add_run('\n\n')
R(p,'選擇依據：聚焦於「如何透過產品創新與行銷策略，解決生成式 AI 對搜尋廣告營收的衝擊」。',size=10)
shading(t1.cell(2,1),'FFF2CC')

# Row3
cell_text(t1.cell(3,0),'3 選擇開放式/封閉式設計',bold=True,size=11); merge(t1,3,1,3)
c=t1.cell(3,1); c.text=''; p=c.paragraphs[0]
for txt,b in [('本研究以',False),('封閉式問卷設計',True),('為主，搭配少量開放式題目：\n\n',False),
    ('封閉式題目（主體）：\n',True),
    ('‧基本資料題：名義尺度、順序尺度\n',False),
    ('‧AI 概覽侵入性：7 點 Likert（Edwards et al., 2002）\n',False),
    ('‧創新廣告格式整合：7 點 Likert（Wojdynski & Evans, 2016）\n',False),
    ('‧外部連結點擊意圖：7 點 Likert（Ajzen, 1991; van Doorn & Hoekstra, 2013）\n',False),
    ('‧搜尋廣告貨幣化效益：7 點 Likert（Rutz & Bucklin, 2011）\n\n',False),
    ('開放式題目（輔助）：\n',True),
    ('‧「請簡述您對 AI 摘要取代傳統搜尋結果的看法」\n',False),
    ('‧「您認為什麼樣的廣告形式最不會干擾您的搜尋體驗？」',False)]:
    R(p,txt,bold=b,size=11)
shading(t1.cell(3,1),'FFF2CC')

# Row4
cell_text(t1.cell(4,0),'4 設計8-10個基本資料與主題相關之行為調查問題',bold=True,size=11); merge(t1,4,1,3)
c=t1.cell(4,1); c.text=''; p=c.paragraphs[0]
for txt,b in [
    ('【人口統計基本資料】\n',True),
    ('Q1. 性別？ □ 男 □ 女 □ 其他 □ 不願透露（名義尺度）\n',False),
    ('Q2. 年齡？ □ 18歲以下 □ 18–24 □ 25–34 □ 35–44 □ 45–54 □ 55以上（順序尺度）\n',False),
    ('Q3. 最高學歷？ □ 高中以下 □ 大專 □ 碩士 □ 博士（順序尺度）\n',False),
    ('Q4. 職業？ □ 學生 □ 上班族 □ 資訊科技 □ 行銷廣告 □ 自營 □ 其他（名義尺度）\n\n',False),
    ('【主題相關行為調查】\n',True),
    ('Q5. 每天使用搜尋引擎頻率？ □ 幾乎不用 □ 1–3次 □ 4–10次 □ 10次以上（順序尺度）\n',False),
    ('Q6. 是否用過 Google AI Overview？ □ 經常 □ 偶爾 □ 知道未用 □ 不知道（順序尺度）\n',False),
    ('Q7. AI 摘要出現時行為？ □ 只看摘要 □ 仍點外部連結 □ 跳過 □ 不確定（名義尺度）\n',False),
    ('Q8. 看到廣告反應？ □ 主動點擊 □ 偶爾 □ 很少 □ 刻意迴避（順序尺度）\n',False),
    ('Q9. 用 AI 聊天工具取代搜尋？ □ 完全取代 □ 部分 □ 偶爾 □ 從未（順序尺度）\n',False),
    ('Q10. 最常用搜尋引擎？（複選）□ Google □ Bing □ Yahoo □ DuckDuckGo □ 其他（名義尺度）',False)]:
    R(p,txt,bold=b,size=10)
shading(t1.cell(4,1),'FFF2CC')

# Row5
merge(t1,5,0,3); cell_text(t1.cell(5,0),'5 列出所選問項與所引用之文獻（需與表2與圖2校準）',bold=True,size=11)

# Row6
for i,h in enumerate(['研究變數','操作型定義','問項 & 尺度','文獻']):
    cell_text(t1.cell(6,i),h,bold=True,size=11); shading(t1.cell(6,i),'FFF2CC')

# Row7: IV1
cell_text(t1.cell(7,0),'AI Overview Intrusiveness\n（AI 概覽侵入性）IV1',bold=True,size=10); shading(t1.cell(7,0),'FFF2CC')
cell_text(t1.cell(7,1),'採用 Edwards et al. (2002) Perceived Intrusiveness Scale，7 點 Likert 量表測量 AI 摘要侵入感受。',size=10); shading(t1.cell(7,1),'FFF2CC')
c=t1.cell(7,2); c.text=''; p=c.paragraphs[0]
for it in ['「當 AI 摘要出現時，我覺得它是…」\n','（7 點 Likert：1＝非常不同意 → 7＝非常同意）\n\n',
    '1. 令人分心的（Distracting）\n','2. 令人不安的（Disturbing）\n','3. 有強迫感的（Forced）\n',
    '4. 干擾性的（Interfering）\n','5. 侵入性的（Intrusive）\n','6. 具有侵犯感的（Invasive）\n',
    '7. 令人突兀的（Obtrusive）\n','α = .90, .85, .88']:
    R(p,it,size=10)
shading(t1.cell(7,2),'FFF2CC')
cell_text(t1.cell(7,3),'Edwards, S. M., Li, H., & Lee, J.-H. (2002). Journal of Advertising, 31(2), 37–47.',size=9); shading(t1.cell(7,3),'FFF2CC')

# Row8: IV2
cell_text(t1.cell(8,0),'Innovative Ad Format Integration\n（創新廣告格式整合）IV2',bold=True,size=10); shading(t1.cell(8,0),'FFF2CC')
cell_text(t1.cell(8,1),'參考 Wojdynski & Evans (2016)，7 點 Likert 量表測量廣告辨識度、態度及來源可信度。',size=10); shading(t1.cell(8,1),'FFF2CC')
c=t1.cell(8,2); c.text=''; p=c.paragraphs[0]
for it in ['「針對嵌入 AI 摘要中的廣告，請評估：」\n','（7 點 Likert）\n\n',
    '▸ 廣告辨識度\n','1. 我能辨識出包含廣告內容。\n','2. 帶有商業推銷目的。\n\n',
    '▸ 可信度\n','3. 相鄰搜尋資訊值得信賴。\n','4. 呈現方式公正客觀。\n\n',
    '▸ 分享意圖\n','5. 願意推薦給朋友。\n','α = .81, .91, .86']:
    R(p,it,size=10)
shading(t1.cell(8,2),'FFF2CC')
cell_text(t1.cell(8,3),'Wojdynski, B. W., & Evans, N. J. (2016). Journal of Advertising, 45(2), 157–168.',size=9); shading(t1.cell(8,3),'FFF2CC')


# ======== 表2 ========
doc.add_paragraph()
p=doc.add_paragraph(); R(p,'針對以上的研究問題，請列出自變數、依變數、以及以上變數的概念型定義與操作型定義。',size=12)
p=doc.add_paragraph(); R(p,'表 2　研究變數說明表',bold=True,size=14)

t2=doc.add_table(rows=5,cols=4); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t2)
for i,h in enumerate(['變數名稱','概念型定義','操作型定義','文獻']):
    cell_text(t2.cell(0,i),h,bold=True,size=11)

for ri,(nm,co,op,rf) in enumerate([
    ('AI Overview Intrusiveness\n（AI 概覽侵入性）\nIV1',
     '使用者感知到 AI 生成摘要干擾其目標導向資訊尋找行為的程度（Li et al., 2002）。',
     '採用 Edwards et al. (2002) 之 Perceived Intrusiveness Scale，7 點 Likert（共 7 題）。',
     'Edwards, S. M., Li, H., & Lee, J.-H. (2002)'),
    ('Innovative Ad Format Integration\n（創新廣告格式整合）\nIV2',
     '搜尋平台在 AI 摘要環境中，將廣告以低辨識度、高內容融合方式嵌入搜尋結果的程度（Wojdynski & Evans, 2016）。',
     '參考 Wojdynski & Evans (2016)，7 點 Likert 量表測量廣告辨識度、態度及來源可信度（共 5 題）。',
     'Wojdynski, B. W., & Evans, N. J. (2016)'),
    ('External URL Click-Through Intention\n（外部連結點擊行為意圖）\nM',
     '使用者在接觸 SERP 後，主動點擊導向第三方外部網站連結的行為意圖強度（Ajzen, 1991）。',
     '以 7 點 Likert 量表，與 Ajzen (1991) TPB 一致，參考 van Doorn & Hoekstra (2013)（共 3 題）。',
     'Ajzen, I. (1991)\nvan Doorn, J., & Hoekstra, J. C. (2013)'),
    ('Search Advertising Monetization Effectiveness\n（搜尋廣告貨幣化效益）\nDV',
     '搜尋平台透過付費關鍵字廣告機制，將使用者搜尋流量變現為廣告收益的效率。',
     '以 7 點 Likert 量表測量品牌注意力、點擊意願與參考價值（共 3 題）。',
     'Rutz, O. J., & Bucklin, R. E. (2011)'),
], start=1):
    cell_text(t2.cell(ri,0),nm,bold=True,size=10,color=(0,86,179)); shading(t2.cell(ri,0),'FBE4D5')
    cell_text(t2.cell(ri,1),co,size=10,color=(0,86,179)); shading(t2.cell(ri,1),'FBE4D5')
    cell_text(t2.cell(ri,2),op,size=10,color=(0,86,179)); shading(t2.cell(ri,2),'FBE4D5')
    cell_text(t2.cell(ri,3),rf,size=9,color=(0,86,179)); shading(t2.cell(ri,3),'FBE4D5')

# ======== 假設表 ========
doc.add_paragraph()
t3=doc.add_table(rows=4,cols=2); t3.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t3)
cell_text(t3.cell(0,0),'假設',bold=True,size=11); cell_text(t3.cell(0,1),'文獻',bold=True,size=11)
for ri,(hy,rf) in enumerate([
    ('H1：AI Overview Intrusiveness 對 External URL Click-Through Intention 具有顯著負向影響',
     'Edwards, S. M., Li, H., & Lee, J.-H. (2002). Journal of Advertising, 31(2), 37–47.'),
    ('H2：External URL Click-Through Intention 對 Search Advertising Monetization Effectiveness 具有顯著正向影響',
     'van Doorn, J., & Hoekstra, J. C. (2013). Marketing Letters, 24(4), 339–351.'),
    ('H3：Innovative Ad Format Integration 對 External URL Click-Through Intention 具有顯著正向影響',
     'Wojdynski, B. W., & Evans, N. J. (2016). Journal of Advertising, 45(2), 157–168.'),
], start=1):
    cell_text(t3.cell(ri,0),hy,size=10,color=(0,86,179)); shading(t3.cell(ri,0),'FBE4D5')
    cell_text(t3.cell(ri,1),rf,size=9,color=(0,86,179)); shading(t3.cell(ri,1),'FBE4D5')
for row in t3.rows: row.cells[0].width=Cm(12); row.cells[1].width=Cm(4)

# ======== 研究架構圖 (editable VML) — AFTER hypothesis table ========
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,'研究架構圖',bold=True,size=14)

# Create editable VML diagram paragraph and insert it BEFORE sectPr
vml_xml = '''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                  xmlns:v="urn:schemas-microsoft-com:vml"
                  xmlns:o="urn:schemas-microsoft-com:office:office">
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:rPr><w:noProof/></w:rPr>
    <w:pict>
      <v:group id="ResearchModel" style="width:480pt;height:220pt" coordorigin="0,0" coordsize="9600,4400">
        <v:rect id="IV1" style="position:absolute;left:0;top:1400;width:2800;height:1800" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>自變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>AI Overview</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Intrusiveness</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;AI &#27010;&#35261;&#20405;&#20837;&#24615;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        <v:rect id="IV2" style="position:absolute;left:3400;top:0;width:2800;height:1200" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>自變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>Innovative Ad Format Integration</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#21109;&#26032;&#24291;&#21578;&#26684;&#24335;&#25972;&#21512;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        <v:rect id="Med" style="position:absolute;left:3400;top:1600;width:2800;height:1600" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>中介變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>External URL</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Click-Through Intention</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#22806;&#37096;&#36899;&#32080;&#40670;&#25802;&#24847;&#22294;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        <v:rect id="DV" style="position:absolute;left:6800;top:1400;width:2800;height:1800" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>依變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Search Advertising</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Monetization Effectiveness</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#25628;&#23563;&#24291;&#21578;&#36008;&#24163;&#21270;&#25928;&#30410;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
        <v:line from="2800,2300" to="3400,2300" strokecolor="black" strokeweight="2pt"><v:stroke endarrow="block"/></v:line>
        <v:rect style="position:absolute;left:2850;top:1850;width:500;height:400" filled="f" stroked="f"><v:textbox inset="0,0,0,0"><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:t>H1</w:t></w:r></w:p></w:txbxContent></v:textbox></v:rect>
        <v:line from="6200,2400" to="6800,2400" strokecolor="black" strokeweight="2pt"><v:stroke endarrow="block"/></v:line>
        <v:rect style="position:absolute;left:6250;top:1950;width:500;height:400" filled="f" stroked="f"><v:textbox inset="0,0,0,0"><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:t>H2</w:t></w:r></w:p></w:txbxContent></v:textbox></v:rect>
        <v:line from="4800,1200" to="4800,1600" strokecolor="black" strokeweight="2pt"><v:stroke endarrow="block"/></v:line>
        <v:rect style="position:absolute;left:4900;top:1200;width:500;height:400" filled="f" stroked="f"><v:textbox inset="0,0,0,0"><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:t>H3</w:t></w:r></w:p></w:txbxContent></v:textbox></v:rect>
      </v:group>
    </w:pict>
  </w:r>
</w:p>'''

diagram_elem = parse_xml(vml_xml)

# Insert before sectPr (which is the last element in body)
sect_pr = doc.element.body.find(qn('w:sectPr'))
if sect_pr is not None:
    sect_pr.addprevious(diagram_elem)
else:
    doc.element.body.append(diagram_elem)


# ======== 報導案例（最末頁） ========
# Insert page break paragraph BEFORE sectPr
pb_p = parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>')
sect_pr = doc.element.body.find(qn('w:sectPr'))
sect_pr.addprevious(pb_p)

# Now add news content paragraphs before sectPr
def add_para_before_sect(text, bold=False, size=12, color=None, bullet=False):
    if bullet:
        xml = f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="ListBullet"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:sz w:val="{size*2}"/></w:rPr><w:t>{text}</w:t></w:r></w:p>'
    else:
        bold_xml = '<w:b/>' if bold else ''
        color_xml = f'<w:color w:val="{color}"/>' if color else ''
        xml = f'<w:p {nsdecls("w")}><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/>{bold_xml}<w:sz w:val="{size*2}"/>{color_xml}</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    elem = parse_xml(xml)
    sect_pr = doc.element.body.find(qn('w:sectPr'))
    sect_pr.addprevious(elem)

add_para_before_sect('報導案例', bold=True, size=16)
add_para_before_sect('Google Gemini 3一戰封神！狠甩ChatGPT，7年來市值首超微軟內幕', bold=True, size=14)
add_para_before_sect('來源：https://www.businessweekly.com.tw/management/blog/3020038', size=10, color='0000FF')
add_para_before_sect('撰文者：王貞懿 編譯 | 商周頭條 2025/11/24', size=10, color='808080')

for b in [
    'Alphabet市值突破3.6兆美元，七年來首度超越微軟，確立AI贏家地位。',
    '透過合併DeepMind與Brain實驗室、創辦人布林回歸督軍，打破穀倉效應。',
    'Google迅速將生成式AI整合至搜尋引擎、YouTube與Android。',
]:
    add_para_before_sect(b, bold=True, size=11)

for text in [
    '華爾街一度熱議Google是否會淪為AI時代的「路殺動物」。2022年底ChatGPT橫空出世，讓搜尋巨頭措手不及。',
    '如今局勢逆轉。Gemini 3在超過二十項基準測試中大幅領先，Alphabet市值突破3.6兆美元，七年來首度超越微軟。',
    '2023年Google合併DeepMind與Brain，全力開發Gemini。執行長皮采打破穀倉、精簡領導層，創辦人布林回歸。',
    'EMARKETER預測Google搜尋廣告市占率明年首次跌破50%。AI Overview出現後，用戶點擊連結比例從15%降至8%。',
    'Google剛用新模型證明了自己，但真正考驗才剛開始：它能領跑多久？代價又是什麼？',
]:
    add_para_before_sect(text, size=11)

# Save
doc.save(OUTPUT)
print(f'✅ Saved: {OUTPUT}')

import os
print(f'File size: {os.path.getsize(OUTPUT):,} bytes')
