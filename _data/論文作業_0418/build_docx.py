"""
Build M11408924_林罕宏_問卷調查設計v2.docx
Single IV: Innovative Ad Format Integration (removed AI Overview Intrusiveness)
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(__file__)
OUTPUT = os.path.join(BASE_DIR, 'M11408924_林罕宏_問卷調查設計v2.docx')

doc = Document()

# ---- Page setup ----
section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Microsoft JhengHei'; style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

# ---- Helpers ----
def shading(cell, color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def cell_text(cell, text, bold=False, size=12, color=None):
    cell.text = ''; p = cell.paragraphs[0]
    r = p.add_run(text); r.font.name='Microsoft JhengHei'; r.font.size=Pt(size); r.font.bold=bold
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    if color: r.font.color.rgb=RGBColor(*color)

def tbl_borders(table):
    table._tbl.tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'))

def merge(table, row, c1, c2):
    table.cell(row, c1).merge(table.cell(row, c2))

def R(para, text, bold=False, size=12, italic=False, color=None):
    r = para.add_run(text); r.bold=bold; r.italic=italic
    r.font.size=Pt(size); r.font.name='Microsoft JhengHei'
    r.element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    if color: r.font.color.rgb=RGBColor(*color)
    return r


# ======== 表1 ========
p=doc.add_paragraph(); R(p,'表 1　問卷調查設計之程序說明表',bold=True,size=14)

# 8 rows: header(研究主題) + step1-4 + step5 header + step5 sub-header + step5 IV row = 8
t1=doc.add_table(rows=8,cols=4); t1.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t1)

# Row0: 研究主題
cell_text(t1.cell(0,0),'研究主題:',bold=True,size=11); merge(t1,0,1,3)
cell_text(t1.cell(0,1),'新型態廣告整合 AI Overview 拯救 Google 廣告衝擊研究',size=11); shading(t1.cell(0,1),'FFF2CC')

# Row1: 步驟1
cell_text(t1.cell(1,0),'1 以80字以內，提出<合作請求>文案',bold=True,size=11); merge(t1,1,1,3)
cell_text(t1.cell(1,1),'您好，我是國立臺北科技大學研究生林罕宏。正在進行「AI 摘要中創新廣告格式對搜尋廣告點擊意願與營收效益之影響」研究，誠摯邀請您撥冗填寫問卷（約 5 分鐘），所有資料僅作學術用途並嚴格保密，感謝您的寶貴協助！',size=11)
shading(t1.cell(1,1),'FFF2CC')

# Row2: 步驟2 — 只保留 Innovative Ad Format Integration
cell_text(t1.cell(2,0),'2 挑選1-2個主要變數',bold=True,size=11); merge(t1,2,1,3)
c=t1.cell(2,1); c.text=''; p=c.paragraphs[0]
R(p,'自變數 IV：Innovative Ad Format Integration（創新廣告格式整合）',bold=True,size=11)
p.add_run('\n\n')
R(p,'選擇依據：管理決策聚焦於「如何透過創新廣告格式設計，在 AI 摘要環境中提升使用者點擊外部連結的意願，進而維持搜尋廣告營收」。本研究以創新廣告格式整合作為核心自變數，探討其對外部連結點擊意圖與搜尋廣告貨幣化效益的影響。',size=10)
shading(t1.cell(2,1),'FFF2CC')

# Row3: 步驟3
cell_text(t1.cell(3,0),'3 選擇開放式/封閉式設計',bold=True,size=11); merge(t1,3,1,3)
c=t1.cell(3,1); c.text=''; p=c.paragraphs[0]
for txt,b in [('本研究以',False),('封閉式問卷設計',True),('為主，搭配少量開放式題目：\n\n',False),
    ('封閉式題目（主體）：\n',True),
    ('‧基本資料題：名義尺度、順序尺度\n',False),
    ('‧創新廣告格式整合：7 點 Likert 量表（Wojdynski & Evans, 2016）\n',False),
    ('‧外部連結點擊意圖：7 點 Likert 量表（Ajzen, 1991; van Doorn & Hoekstra, 2013）\n',False),
    ('‧搜尋廣告貨幣化效益：7 點 Likert 量表（Rutz & Bucklin, 2011）\n\n',False),
    ('開放式題目（輔助）：\n',True),
    ('‧「請簡述您對 AI 摘要取代傳統搜尋結果的看法」\n',False),
    ('‧「您認為什麼樣的廣告形式最不會干擾您的搜尋體驗？」',False)]:
    R(p,txt,bold=b,size=11)
shading(t1.cell(3,1),'FFF2CC')

# Row4: 步驟4
cell_text(t1.cell(4,0),'4 設計8-10個基本資料與主題相關之行為調查問題',bold=True,size=11); merge(t1,4,1,3)
c=t1.cell(4,1); c.text=''; p=c.paragraphs[0]
for txt,b in [
    ('【人口統計基本資料】\n',True),
    ('Q1. 性別？ □ 男 □ 女 □ 其他 □ 不願透露（名義尺度）\n',False),
    ('Q2. 年齡？ □ 18歲以下 □ 18–24 □ 25–34 □ 35–44 □ 45–54 □ 55以上（順序尺度）\n',False),
    ('Q3. 最高學歷？ □ 高中以下 □ 大專 □ 碩士 □ 博士（順序尺度）\n',False),
    ('Q4. 職業？ □ 學生 □ 上班族 □ 資訊科技 □ 行銷廣告 □ 自營 □ 其他（名義尺度）\n\n',False),
    ('【主題相關行為調查】\n',True),
    ('Q5. 每天使用搜尋引擎頻率？ □ 幾乎不用 □ 1–3次 □ 4–10次 □ 10次以上（順序尺度）\n',False),
    ('Q6. 是否用過 Google AI Overview？ □ 經常 □ 偶爾 □ 知道未用 □ 不知道（順序尺度）\n',False),
    ('Q7. AI 摘要出現時行為？ □ 只看摘要 □ 仍點外部連結 □ 跳過 □ 不確定（名義尺度）\n',False),
    ('Q8. 看到廣告反應？ □ 主動點擊 □ 偶爾 □ 很少 □ 刻意迴避（順序尺度）\n',False),
    ('Q9. 用 AI 聊天工具取代搜尋？ □ 完全取代 □ 部分 □ 偶爾 □ 從未（順序尺度）\n',False),
    ('Q10. 最常用搜尋引擎？（複選）□ Google □ Bing □ Yahoo □ DuckDuckGo □ 其他（名義尺度）',False)]:
    R(p,txt,bold=b,size=10)
shading(t1.cell(4,1),'FFF2CC')

# Row5: 步驟5 header
merge(t1,5,0,3); cell_text(t1.cell(5,0),'5 列出所選問項與所引用之文獻（需與表2與圖2校準）',bold=True,size=11)

# Row6: 子表 header
for i,h in enumerate(['研究變數','操作型定義','問項 & 尺度','文獻']):
    cell_text(t1.cell(6,i),h,bold=True,size=11); shading(t1.cell(6,i),'FFF2CC')

# Row7: IV (only Innovative Ad Format Integration)
cell_text(t1.cell(7,0),'Innovative Ad Format Integration\n（創新廣告格式整合）IV',bold=True,size=10); shading(t1.cell(7,0),'FFF2CC')
cell_text(t1.cell(7,1),'參考 Wojdynski & Evans (2016) 之實驗設計，操弄廣告揭露位置（頂部 vs. 中段 vs. 底部）與揭露語言（「廣告」vs.「贊助內容」vs.「品牌聲音」），以 7 點 Likert 量表測量受試者對廣告的辨識度、廣告態度及來源可信度。',size=10); shading(t1.cell(7,1),'FFF2CC')
c=t1.cell(7,2); c.text=''; p=c.paragraphs[0]
for it in ['「針對嵌入 AI 摘要中的廣告內容，請評估您的感受：」\n',
    '（7 點 Likert：1＝非常不同意 → 7＝非常同意）\n\n',
    '▸ 廣告辨識度（Ad Recognition）\n',
    '1. 我能辨識出該頁面中包含廣告內容。\n',
    '2. 我覺得該內容帶有商業推銷目的。\n\n',
    '▸ 新聞可信度（Perceived News Credibility）\n',
    '3. 與廣告相鄰的搜尋資訊是值得信賴的。\n',
    '4. 該廣告呈現方式讓我覺得公正、客觀。\n\n',
    '▸ 分享意圖（Intention to Share）\n',
    '5. 我願意推薦這個搜尋結果給朋友。\n',
    'α = .81, .91, .86']:
    R(p,it,size=10)
shading(t1.cell(7,2),'FFF2CC')
cell_text(t1.cell(7,3),'Wojdynski, B. W., & Evans, N. J. (2016). Going Native: Effects of Disclosure Position and Language on the Recognition and Evaluation of Online Native Advertising. Journal of Advertising, 45(2), 157–168.',size=9); shading(t1.cell(7,3),'FFF2CC')


# ======== 表2（3 變數：IV, M, DV） ========
doc.add_paragraph()
p=doc.add_paragraph(); R(p,'針對以上的研究問題，請列出自變數、依變數、以及以上變數的概念型定義與操作型定義。',size=12)
p=doc.add_paragraph(); R(p,'表 2　研究變數說明表',bold=True,size=14)

t2=doc.add_table(rows=4,cols=4); t2.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t2)
for i,h in enumerate(['變數名稱','概念型定義','操作型定義','文獻']):
    cell_text(t2.cell(0,i),h,bold=True,size=11)

for ri,(nm,co,op,rf) in enumerate([
    ('Innovative Ad Format Integration\n（創新廣告格式整合）\nIV',
     '原生廣告（native advertising）被定義為一種與周圍編輯內容在視覺外觀與功能上高度一致的付費媒體形式（Wojdynski & Evans, 2016）。本研究將創新廣告格式整合定義為：搜尋平台在 AI 摘要環境中，透過揭露位置（disclosure position）與揭露語言（disclosure language）的設計，將廣告以低辨識度、高內容融合的方式嵌入搜尋結果的程度。',
     '參考 Wojdynski & Evans (2016) 之實驗設計，操弄廣告揭露位置（頂部 vs. 中段 vs. 底部）與揭露語言，以 7 點 Likert 量表測量受試者對廣告的辨識度（ad recognition）、廣告態度（ad evaluation）及來源可信度（source credibility）（共 5 題）。',
     'Wojdynski, B. W., & Evans, N. J. (2016)'),
    ('External URL Click-Through Intention\n（外部連結點擊行為意圖）\nM',
     '依據計畫行為理論（TPB），行為意圖係個體對特定行為之主觀可能性評估（Ajzen, 1991）。本研究將外部連結點擊意圖定義為：使用者在接觸搜尋結果頁面（SERP）後，主動點擊導向第三方外部網站連結的行為意圖強度。',
     '以 7 點 Likert 量表（「我願意點擊連結」「我打算點擊連結」「我可能會去點擊連結」），與 Ajzen (1991) TPB 框架之行為意圖測量方式一致，參考 van Doorn & Hoekstra (2013) 對點擊意圖的操作化設計（共 3 題）。',
     'Ajzen, I. (1991)\nvan Doorn, J., & Hoekstra, J. C. (2013)'),
    ('Search Advertising Monetization Effectiveness\n（搜尋廣告貨幣化效益）\nDV',
     '搜尋廣告貨幣化效益定義為：搜尋平台透過付費關鍵字廣告（paid search advertising）機制，將使用者搜尋流量變現為廣告收益的效率，以 SERP 廣告點擊份額（click share）、每次點擊成本（CPC）及廣告曝光轉換率（CTR）綜合衡量。',
     '以 7 點 Likert 量表測量使用者對 AI 摘要中廣告之品牌注意力、點擊意願與參考價值（共 3 題）。',
     'Rutz, O. J., & Bucklin, R. E. (2011)'),
], start=1):
    cell_text(t2.cell(ri,0),nm,bold=True,size=10,color=(0,86,179)); shading(t2.cell(ri,0),'FBE4D5')
    cell_text(t2.cell(ri,1),co,size=10,color=(0,86,179)); shading(t2.cell(ri,1),'FBE4D5')
    cell_text(t2.cell(ri,2),op,size=10,color=(0,86,179)); shading(t2.cell(ri,2),'FBE4D5')
    cell_text(t2.cell(ri,3),rf,size=9,color=(0,86,179)); shading(t2.cell(ri,3),'FBE4D5')

# ======== 假設表（2 假設：H1, H2） ========
doc.add_paragraph()
t3=doc.add_table(rows=3,cols=2); t3.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl_borders(t3)
cell_text(t3.cell(0,0),'假設',bold=True,size=11); cell_text(t3.cell(0,1),'文獻',bold=True,size=11)
for ri,(hy,rf) in enumerate([
    ('H1：Innovative Ad Format Integration（創新廣告格式整合）對 External URL Click-Through Intention（外部連結點擊意圖）具有顯著正向影響',
     'Wojdynski, B. W., & Evans, N. J. (2016). Going Native: Effects of Disclosure Position and Language on the Recognition and Evaluation of Online Native Advertising. Journal of Advertising, 45(2), 157–168.'),
    ('H2：External URL Click-Through Intention（外部連結點擊意圖）對 Search Advertising Monetization Effectiveness（搜尋廣告貨幣化效益）具有顯著正向影響',
     'van Doorn, J., & Hoekstra, J. C. (2013). Customization of Online Advertising: The Role of Intrusiveness. Marketing Letters, 24(4), 339–351.'),
], start=1):
    cell_text(t3.cell(ri,0),hy,size=10,color=(0,86,179)); shading(t3.cell(ri,0),'FBE4D5')
    cell_text(t3.cell(ri,1),rf,size=9,color=(0,86,179)); shading(t3.cell(ri,1),'FBE4D5')
for row in t3.rows: row.cells[0].width=Cm(12); row.cells[1].width=Cm(4)

# ======== 研究架構圖 (editable VML) — IV → M → DV 線性模型 ========
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,'研究架構圖',bold=True,size=14)

# Simple linear: IV → M → DV
vml_xml = '''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                  xmlns:v="urn:schemas-microsoft-com:vml"
                  xmlns:o="urn:schemas-microsoft-com:office:office">
  <w:pPr><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:rPr><w:noProof/></w:rPr>
    <w:pict>
      <v:group id="ResearchModelV2" style="width:480pt;height:140pt" coordorigin="0,0" coordsize="9600,2800">

        <!-- IV Box (left) -->
        <v:rect id="IV" style="position:absolute;left:0;top:400;width:2800;height:1800" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>自變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Innovative Ad Format</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Integration</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#21109;&#26032;&#24291;&#21578;&#26684;&#24335;&#25972;&#21512;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>

        <!-- Mediator Box (center) -->
        <v:rect id="Med" style="position:absolute;left:3400;top:400;width:2800;height:1800" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>中介變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>External URL</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Click-Through Intention</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#22806;&#37096;&#36899;&#32080;&#40670;&#25802;&#24847;&#22294;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>

        <!-- DV Box (right) -->
        <v:rect id="DV" style="position:absolute;left:6800;top:400;width:2800;height:1800" fillcolor="white" strokecolor="black" strokeweight="2pt">
          <v:textbox inset="4pt,4pt,4pt,4pt">
            <w:txbxContent>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:b/><w:color w:val="0056B3"/><w:sz w:val="24"/><w:u w:val="single"/></w:rPr><w:t>依變數</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Search Advertising</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="20"/></w:rPr><w:t>Monetization Effectiveness</w:t></w:r></w:p>
              <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/><w:color w:val="0056B3"/><w:sz w:val="18"/></w:rPr><w:t>&#65288;&#25628;&#23563;&#24291;&#21578;&#36008;&#24163;&#21270;&#25928;&#30410;&#65289;</w:t></w:r></w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>

        <!-- Arrow H1: IV → M -->
        <v:line from="2800,1300" to="3400,1300" strokecolor="black" strokeweight="2pt"><v:stroke endarrow="block"/></v:line>
        <v:rect style="position:absolute;left:2850;top:800;width:500;height:400" filled="f" stroked="f">
          <v:textbox inset="0,0,0,0"><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:t>H1</w:t></w:r></w:p></w:txbxContent></v:textbox>
        </v:rect>

        <!-- Arrow H2: M → DV -->
        <v:line from="6200,1300" to="6800,1300" strokecolor="black" strokeweight="2pt"><v:stroke endarrow="block"/></v:line>
        <v:rect style="position:absolute;left:6250;top:800;width:500;height:400" filled="f" stroked="f">
          <v:textbox inset="0,0,0,0"><w:txbxContent><w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="22"/></w:rPr><w:t>H2</w:t></w:r></w:p></w:txbxContent></v:textbox>
        </v:rect>

      </v:group>
    </w:pict>
  </w:r>
</w:p>'''

diagram_elem = parse_xml(vml_xml)
sect_pr = doc.element.body.find(qn('w:sectPr'))
sect_pr.addprevious(diagram_elem)


# ======== 報導案例（最末頁） ========
def add_before_sect(text, bold=False, size=12, color=None):
    bold_xml = '<w:b/>' if bold else ''
    color_xml = f'<w:color w:val="{color}"/>' if color else ''
    xml = f'<w:p {nsdecls("w")}><w:r><w:rPr><w:rFonts w:eastAsia="Microsoft JhengHei"/>{bold_xml}<w:sz w:val="{size*2}"/>{color_xml}</w:rPr><w:t xml:space="preserve">{text}</w:t></w:r></w:p>'
    doc.element.body.find(qn('w:sectPr')).addprevious(parse_xml(xml))

# Page break
doc.element.body.find(qn('w:sectPr')).addprevious(
    parse_xml(f'<w:p {nsdecls("w")}><w:r><w:br w:type="page"/></w:r></w:p>'))

add_before_sect('報導案例', bold=True, size=16)
add_before_sect('Google Gemini 3一戰封神！狠甩ChatGPT，7年來市值首超微軟內幕', bold=True, size=14)
add_before_sect('來源：https://www.businessweekly.com.tw/management/blog/3020038', size=10, color='0000FF')
add_before_sect('撰文者：王貞懿 編譯 | 商周頭條 2025/11/24', size=10, color='808080')

for t in [
    'Alphabet市值突破3.6兆美元，七年來首度超越微軟，確立AI贏家地位。',
    '透過合併DeepMind與Brain實驗室、創辦人布林回歸督軍，打破穀倉效應。',
    'Google迅速將生成式AI整合至搜尋引擎、YouTube與Android。',
]:
    add_before_sect(t, bold=True, size=11)

for t in [
    '華爾街一度熱議Google是否會淪為AI時代的「路殺動物」。2022年底ChatGPT橫空出世，讓搜尋巨頭措手不及。',
    '如今局勢逆轉。Gemini 3在超過二十項基準測試中大幅領先，Alphabet市值突破3.6兆美元，七年來首度超越微軟。',
    '2023年Google合併DeepMind與Brain，全力開發Gemini。執行長皮采打破穀倉、精簡領導層，創辦人布林回歸。',
    'EMARKETER預測Google搜尋廣告市占率明年首次跌破50%。AI Overview出現後，用戶點擊連結比例從15%降至8%。',
    'Google剛用新模型證明了自己，但真正考驗才剛開始：它能領跑多久？代價又是什麼？',
]:
    add_before_sect(t, size=11)

# Save
doc.save(OUTPUT)
print(f'✅ Saved: {OUTPUT}')
print(f'File size: {os.path.getsize(OUTPUT):,} bytes')
